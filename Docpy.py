import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL

def Crear_Documento_Word_Cuadricula_Expandida(
    DataFrame: pd.DataFrame, 
    Columna_Normal: str, 
    Columna_Moneda: str, 
    Archivo_Salida: str, 
    Filas_Por_Pagina: int, 
    Columnas_Por_Pagina: int,
    Longitud_Pagina: int = 22
) -> None:

    """
    Crea un documento Word con un diseño de cuadrícula, donde las filas 
    y columnas se expanden para cubrir toda la página. La columna 
    numérica es negrita y ocupa la mayoría de la celda, mientras que la 
    columna de texto es más pequeña.

    Parámetros:
        DataFrame (pd.DataFrame): El dataframe fuente que contiene 
        los datos.
        Columna_Normal (str): El nombre de la columna cuyos valores 
        aparecerán más pequeños y sin negrita.
        Columna_Moneda (str): El nombre de la columna cuyos valores 
        aparecerán en negrita y más grandes.
        Archivo_Salida (str): El nombre del archivo Word de salida.
        Filas_Por_Pagina (int): Número de filas a mostrar por página.
        Columnas_Por_Pagina (int): Número de columnas a mostrar por fila.
        Longitud_Pagina (int): Centímetros de la página.

    Retorna:
        None

    Ejemplo:
        >>> df = pd.DataFrame({'Nombre': ['Alicia', 'Bob', 'Carlos'], 
        ...                    'Monto': [1234.56, 789.01, 456.78]})
        >>> Crear_Documento_Word_Cuadricula_Expandida(
        ...     df, 'Nombre', 'Monto', 'salida.docx', 5, 2, 22
        ... )

    """

    Archivo_Documento = Document()
    Total_Entradas = len(DataFrame)
    Entradas_Por_Pagina = Filas_Por_Pagina * Columnas_Por_Pagina

    Indice_Entrada = 0

    while Indice_Entrada < Total_Entradas:
        # Agregar una nueva sección (página).
        if Indice_Entrada > 0:
            Archivo_Documento.add_page_break()

        # Crear una tabla para la página.
        Tabla = Archivo_Documento.add_table(
            rows=Filas_Por_Pagina, 
            cols=Columnas_Por_Pagina
        )

        # Ajustar la tabla para expandirse a través de toda la página.
        Ancho_Tabla = 16.5  # Ancho en cm (ajustar según tamaño y 
                           # márgenes).
        Ancho_Columna = Ancho_Tabla / Columnas_Por_Pagina

        for Indice_Fila in range(Filas_Por_Pagina):
            for Indice_Col in range(Columnas_Por_Pagina):
                Celda = Tabla.cell(Indice_Fila, Indice_Col)

                # Establecer ancho de celda.
                Elemento_Ancho_Celda = Celda._element
                Propiedades_Tabla = (
                    Elemento_Ancho_Celda.xpath('.//w:tcPr')[0]
                )
                Elemento_Ancho = OxmlElement('w:tcW')
                Elemento_Ancho.set(
                    qn('w:w'), 
                    str(int(Ancho_Columna * 567))
                )  # Convertir cm a twips.
                Elemento_Ancho.set(qn('w:type'), 'dxa')
                Propiedades_Tabla.append(Elemento_Ancho)

                if Indice_Entrada < Total_Entradas:
                    # Obtener la fila actual del DataFrame.
                    Fila = DataFrame.iloc[Indice_Entrada]
                    Texto_Pequeño = Fila[Columna_Normal]
                    Valor_Grande_Negrita = (
                        f"${int(Fila[Columna_Moneda]):,}"
                    )

                    # Agregar contenido a la celda.
                    Parrafo = Celda.paragraphs[0]

                    Ejecucion_Negrita = Parrafo.add_run(
                        Valor_Grande_Negrita
                    )
                    Ejecucion_Negrita.bold = True
                    
                    Ejecucion_Normal = Parrafo.add_run(
                        f"\n{Texto_Pequeño}"
                    )

                    # Ajustar tamaño de fuente para el nuevo diseño.
                    Tamaño_Fuente_Grande = 48  # Tamaño grande para 
                                              # números en negrita.
                    Tamaño_Fuente_Pequeña = 12  # Tamaño pequeño para 
                                               # texto normal.
                    Ejecucion_Negrita.font.size = Pt(
                        Tamaño_Fuente_Grande
                    )
                    Ejecucion_Normal.font.size = Pt(
                        Tamaño_Fuente_Pequeña
                    )

                    # Alinear el contenido en el centro de la celda.
                    Parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    Celda.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    Indice_Entrada += 1
                else:
                    # Dejar celdas vacías si no hay más entradas.
                    Celda.text = ""

        # Ajustar altura de fila para expandir filas a través de la 
        # página.
        for Fila in Tabla.rows:
            Altura_Fila = Longitud_Pagina / Filas_Por_Pagina
            Fila.height = Cm(Altura_Fila)

        # Aplicar un estilo de cuadrícula para visibilidad.
        Tabla.style = "Table Grid"

    # Guardar el documento en el archivo especificado.
    Archivo_Documento.save(Archivo_Salida)